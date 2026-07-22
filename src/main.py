from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.middleware.proxy_fix import ProxyFix
import io
import json
import os
import re
import logging
import threading
import time
import requests
from collections import OrderedDict
from datetime import datetime
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import (
    NoTranscriptFound,
    TranscriptsDisabled,
    VideoUnavailable,
    RequestBlocked,
    IpBlocked,
)
from youtube_transcript_api.formatters import SRTFormatter, WebVTTFormatter
from youtube_transcript_api.proxies import WebshareProxyConfig, GenericProxyConfig

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
# Trust Render's reverse proxy so rate limiting sees the real client IP
app.wsgi_app = ProxyFix(app.wsgi_app, x_for=1, x_proto=1)

limiter = Limiter(
    get_remote_address,
    app=app,
    storage_uri='memory://',
    headers_enabled=True,
)

ALLOWED_ORIGINS = [
    'https://transcriptflow.io',
    'https://www.transcriptflow.io',
    'https://transcriptflow-frontend.pages.dev',
    'http://localhost:5173',
    'http://localhost:4173',
    'http://localhost:3000',
]
CORS(app, origins=ALLOWED_ORIGINS, methods=['GET', 'POST', 'OPTIONS'], allow_headers=['Content-Type', 'Authorization'])


class RotatingGenericProxyConfig(GenericProxyConfig):
    """GenericProxyConfig with retries: each retry goes through a fresh rotated IP,
    so a single blocked exit IP doesn't fail the whole request."""

    @property
    def retries_when_blocked(self):
        return int(os.environ.get('PROXY_RETRIES_WHEN_BLOCKED', 4))


def _proxy_config_from_env():
    """Build an optional proxy config from environment variables.

    WEBSHARE_PROXY_USERNAME / WEBSHARE_PROXY_PASSWORD -> Webshare rotating residential proxy
    GENERIC_PROXY_URL -> any http(s) proxy (e.g. a self-hosted residential proxy)
    Neither set -> direct connection.
    """
    webshare_user = os.environ.get('WEBSHARE_PROXY_USERNAME')
    webshare_pass = os.environ.get('WEBSHARE_PROXY_PASSWORD')
    if webshare_user and webshare_pass:
        logger.info("Using Webshare proxy configuration")
        return WebshareProxyConfig(proxy_username=webshare_user, proxy_password=webshare_pass)

    generic_url = os.environ.get('GENERIC_PROXY_URL')
    if generic_url:
        logger.info("Using generic proxy configuration")
        return RotatingGenericProxyConfig(http_url=generic_url, https_url=generic_url)

    return None


PROXY_CONFIG = _proxy_config_from_env()

# In-memory LRU transcript cache. Repeat requests for the same video are
# served from here instead of re-fetching through the (paid) proxy.
CACHE_MAX_ENTRIES = int(os.environ.get('TRANSCRIPT_CACHE_MAX', 500))
CACHE_TTL_SECONDS = int(os.environ.get('TRANSCRIPT_CACHE_TTL', 24 * 3600))
_cache = OrderedDict()
_cache_lock = threading.Lock()


def cache_get(video_id):
    with _cache_lock:
        entry = _cache.get(video_id)
        if entry is None:
            return None
        stored_at, payload = entry
        if time.time() - stored_at > CACHE_TTL_SECONDS:
            del _cache[video_id]
            return None
        _cache.move_to_end(video_id)
        return dict(payload)


def cache_set(video_id, payload):
    with _cache_lock:
        _cache[video_id] = (time.time(), dict(payload))
        _cache.move_to_end(video_id)
        while len(_cache) > CACHE_MAX_ENTRIES:
            _cache.popitem(last=False)


VIDEO_ID_RE = re.compile(r'(?:v=|youtu\.be/|embed/|watch\?v=)([a-zA-Z0-9_-]{11})')


def _video_id_from_url(video_url):
    match = VIDEO_ID_RE.search(video_url or '')
    return match.group(1) if match else None


def _get_video_title(video_id):
    """Best-effort video title via YouTube's public oEmbed endpoint (not IP-blocked)."""
    try:
        resp = requests.get(
            'https://www.youtube.com/oembed',
            params={'url': f'https://www.youtube.com/watch?v={video_id}', 'format': 'json'},
            timeout=5,
        )
        if resp.ok:
            return resp.json().get('title')
    except Exception:
        pass
    return None


def _pick_transcript(transcript_list):
    """Prefer a manually created transcript, then English, then whatever exists."""
    available = list(transcript_list)
    if not available:
        return None
    manual = [t for t in available if not t.is_generated]
    for pool in (manual, available):
        for t in pool:
            if t.language_code.startswith('en'):
                return t
    return (manual or available)[0]


class TranslationUnavailable(Exception):
    pass


def _fetch_transcript(video_id, target_language=None):
    """Fetch a transcript, optionally translated to target_language.

    Returns (fetched_transcript, language_code).
    """
    ytt_api = YouTubeTranscriptApi(proxy_config=PROXY_CONFIG)

    if target_language:
        transcript_list = ytt_api.list(video_id)
        transcript = _pick_transcript(transcript_list)
        if transcript is None:
            raise NoTranscriptFound(video_id, [], None)
        if transcript.language_code == target_language or transcript.language_code.split('-')[0] == target_language:
            fetched = transcript.fetch()
            return fetched, transcript.language_code
        if not transcript.is_translatable:
            raise TranslationUnavailable()
        codes = {tl.language_code for tl in transcript.translation_languages}
        if target_language not in codes:
            raise TranslationUnavailable()
        fetched = transcript.translate(target_language).fetch()
        return fetched, target_language

    try:
        fetched = ytt_api.fetch(video_id)
        return fetched, fetched.language_code
    except NoTranscriptFound:
        transcript_list = ytt_api.list(video_id)
        transcript = _pick_transcript(transcript_list)
        if transcript is None:
            raise NoTranscriptFound(video_id, [], None)
        fetched = transcript.fetch()
        return fetched, transcript.language_code


# Exceptions that mean "this video has no usable transcript" — never retry these.
_PERMANENT_ERRORS = (NoTranscriptFound, TranscriptsDisabled, VideoUnavailable, TranslationUnavailable)


def _fetch_with_retries(video_id, target_language=None, attempts=3):
    """Retry transient failures (blocked/captcha'd rotated proxy IPs get a fresh
    exit IP on each attempt); permanent no-transcript errors are raised immediately."""
    last_exc = None
    for attempt in range(attempts):
        try:
            return _fetch_transcript(video_id, target_language)
        except _PERMANENT_ERRORS:
            raise
        except Exception as e:
            last_exc = e
            logger.warning(f"Transient fetch failure for video {video_id} (attempt {attempt + 1}/{attempts}): {type(e).__name__}")
    raise last_exc


def _transcript_payload(video_id, target_language=None):
    """Return (payload, was_cached) for a video, serving from the LRU cache when
    possible. Raises transcript errors on failure."""
    cache_key = f"{video_id}:{target_language or 'default'}"
    cached = cache_get(cache_key)
    if cached is not None:
        return cached, True

    fetched_transcript, language_code = _fetch_with_retries(video_id, target_language)
    transcript_data = fetched_transcript.to_raw_data()

    formatted_transcript = ""
    for entry in transcript_data:
        start_time_seconds = int(entry['start'])
        minutes = start_time_seconds // 60
        seconds = start_time_seconds % 60
        formatted_transcript += f"[{minutes:02d}:{seconds:02d}] {entry['text']}\n"

    srt_content = SRTFormatter().format_transcript(fetched_transcript)
    vtt_content = WebVTTFormatter().format_transcript(fetched_transcript)
    word_count = sum(len(entry['text'].split()) for entry in transcript_data)
    last_entry = transcript_data[-1] if transcript_data else {'start': 0, 'duration': 0}
    duration_seconds = int(last_entry['start'] + last_entry.get('duration', 0))
    video_title = _get_video_title(video_id)

    payload = {
        'transcript': formatted_transcript.strip(),
        'language': language_code,
        'video_id': video_id,
        'video_title': video_title,
        'word_count': word_count,
        'duration': duration_seconds,
        'srt': srt_content,
        'vtt': vtt_content,
        'cached': False,
        'success': True
    }
    cache_set(cache_key, payload)
    _increment_transcript_counter()
    logger.info(f"Fetched and cached transcript for video {video_id} ({target_language or 'default'})")
    return payload, False


def _transcript_error_response(exc, video_id):
    """Map transcript exceptions to (json, status) responses; None if unhandled."""
    if isinstance(exc, TranslationUnavailable):
        return jsonify({
            'error': 'This transcript cannot be translated to the requested language.',
            'error_type': 'translation_unavailable'
        }), 400
    if isinstance(exc, NoTranscriptFound):
        return jsonify({
            'error': 'No transcript available for this video. The video may not have captions or subtitles.',
            'error_type': 'no_transcript'
        }), 404
    if isinstance(exc, VideoUnavailable):
        return jsonify({
            'error': 'Video is not accessible. It may be private, deleted, or restricted.',
            'error_type': 'video_unavailable'
        }), 404
    if isinstance(exc, TranscriptsDisabled):
        return jsonify({
            'error': 'Transcripts are disabled for this video.',
            'error_type': 'transcripts_disabled'
        }), 404
    if isinstance(exc, (RequestBlocked, IpBlocked)):
        logger.error(f"YouTube blocked the request for video {video_id} (server IP is blocked)")
        return jsonify({
            'error': 'YouTube is temporarily blocking requests from our server. Please try again in a few minutes.',
            'error_type': 'ip_blocked'
        }), 503
    return None


def _slugify(text):
    slug = re.sub(r'[^a-z0-9]+', '-', (text or '').lower()).strip('-')[:60]
    return slug or 'transcript'


FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'fonts')
TIMESTAMP_LINE_RE = re.compile(r'^(\[\d+:\d{2}\])\s?(.*)$')


def _build_docx(payload):
    """Build a Word document of the transcript in memory."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(payload.get('video_title') or 'YouTube Video Transcript', level=1)
    meta = doc.add_paragraph()
    meta_run = meta.add_run(
        f"Language: {payload.get('language')}   •   Words: {payload.get('word_count')}   •   "
        f"Source: https://youtu.be/{payload.get('video_id')}   •   Generated by transcriptflow.io"
    )
    meta_run.font.size = Pt(9)

    for line in (payload.get('transcript') or '').split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        match = TIMESTAMP_LINE_RE.match(line)
        if match:
            ts_run = p.add_run(match.group(1) + ' ')
            ts_run.bold = True
            p.add_run(match.group(2))
        else:
            p.add_run(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf


# Scripts that need extra fonts / text shaping in the PDF
_DEVANAGARI_RE = re.compile(r'[ऀ-ॿ]')
_ARABIC_RE = re.compile(r'[؀-ۿݐ-ݿࢠ-ࣿ]')
_COMPLEX_RE = re.compile(r'[֐-ࣿऀ-෿฀-໿ក-៿]')

# fontTools logs every subsetting step at INFO — silence it
logging.getLogger('fontTools').setLevel(logging.WARNING)


def _build_pdf(payload):
    """Single-video PDF via the Typst compiler.

    Typst shapes complex scripts (Devanagari, Arabic) dramatically faster than
    fpdf2's HarfBuzz path — the difference between seconds and minutes on a long
    (multi-hour) transcript. Falls back to the fpdf2 builder on any failure.
    """
    try:
        title = payload.get('video_title') or 'YouTube Video Transcript'
        meta = (
            f"Language: {payload.get('language')}  •  Words: {payload.get('word_count')}  •  "
            f"Source: https://youtu.be/{payload.get('video_id')}  •  Generated by transcriptflow.io"
        )
        parts = [
            '#set text(font: ("Noto Sans", "Noto Sans Devanagari", "Noto Sans Arabic"), size: 10.5pt)',
            '#set page(margin: 18mm)',
            '#set par(leading: 0.45em, spacing: 0.5em)',
            '',
            f'#text(size: 15pt, weight: "bold")[{_typst_escape(title)}]',
            '',
            f'#text(size: 9pt, fill: rgb("#787878"))[{_typst_escape(meta)}]',
            '',
        ]
        body = '\n\n'.join(
            _typst_escape(line) for line in (payload.get('transcript') or '').split('\n') if line.strip()
        )
        parts.append(body)
        return _typst_compile('\n'.join(parts))
    except Exception as e:
        logger.warning(f"Typst single PDF build failed ({type(e).__name__}: {e}); falling back to fpdf2")
        return _build_pdf_fpdf(payload)


def _build_pdf_fpdf(payload):
    """Fallback single-video PDF builder using fpdf2 (slow for complex scripts at scale).

    Latin fast path (2 fonts, no shaping); extra fonts + HarfBuzz shaping only
    when the text actually needs them.
    """
    from fpdf import FPDF

    text_sample = f"{payload.get('video_title') or ''}\n{payload.get('transcript') or ''}"
    pdf = FPDF()
    pdf.add_font('NotoSans', '', os.path.join(FONT_DIR, 'NotoSans-Regular.ttf'))
    pdf.add_font('NotoSans', 'B', os.path.join(FONT_DIR, 'NotoSans-Bold.ttf'))

    fallbacks = []
    if _DEVANAGARI_RE.search(text_sample):
        pdf.add_font('NotoDeva', '', os.path.join(FONT_DIR, 'NotoSansDevanagari-Regular.ttf'))
        fallbacks.append('NotoDeva')
    if _ARABIC_RE.search(text_sample):
        pdf.add_font('NotoArab', '', os.path.join(FONT_DIR, 'NotoSansArabic-Regular.ttf'))
        fallbacks.append('NotoArab')
    if fallbacks:
        pdf.set_fallback_fonts(fallbacks)
    if _COMPLEX_RE.search(text_sample):
        try:
            # Proper shaping/BiDi for Arabic and Indic scripts (needs uharfbuzz)
            pdf.set_text_shaping(True)
        except Exception:
            pass

    pdf.set_auto_page_break(True, margin=18)
    pdf.add_page()
    pdf.set_font('NotoSans', 'B', 15)
    pdf.multi_cell(0, 8, payload.get('video_title') or 'YouTube Video Transcript', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(1)
    pdf.set_font('NotoSans', '', 9)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(
        0, 5.5,
        f"Language: {payload.get('language')}   •   Words: {payload.get('word_count')}   •   "
        f"Source: https://youtu.be/{payload.get('video_id')}   •   Generated by transcriptflow.io",
        new_x='LMARGIN', new_y='NEXT'
    )
    pdf.ln(3)
    pdf.set_text_color(20, 20, 20)
    pdf.set_font('NotoSans', '', 10.5)
    for line in (payload.get('transcript') or '').split('\n'):
        pdf.multi_cell(0, 6, line, new_x='LMARGIN', new_y='NEXT')

    return io.BytesIO(bytes(pdf.output()))


EXPORT_FORMATS = {
    'txt': ('text/plain', None),
    'docx': ('application/vnd.openxmlformats-officedocument.wordprocessingml.document', _build_docx),
    'pdf': ('application/pdf', _build_pdf),
}


def _typst_escape(text):
    """Escape characters that are markup in Typst so user text renders literally."""
    if not text:
        return ''
    # Backslash first, then the rest.
    text = text.replace('\\', '\\\\')
    for ch in ('#', '*', '_', '`', '$', '<', '>', '@', '~', '"', '='):
        text = text.replace(ch, '\\' + ch)
    return text


def _typst_compile(source):
    """Compile a Typst source string to PDF bytes, using our bundled fonts."""
    import typst
    import tempfile
    with tempfile.TemporaryDirectory() as tmp:
        src_path = os.path.join(tmp, 'doc.typ')
        with open(src_path, 'w', encoding='utf-8') as fh:
            fh.write(source)
        return io.BytesIO(typst.compile(src_path, font_paths=[FONT_DIR]))


def _build_combined_pdf(playlist_title, videos):
    """Build a single PDF of every video's transcript via the Typst compiler.

    Typst shapes complex scripts (Devanagari, Arabic) ~50x faster than fpdf2's
    HarfBuzz path, which is the difference between seconds and many minutes on a
    large multi-video playlist. Falls back to the fpdf2 builder on any failure.
    """
    try:
        parts = [
            '#set text(font: ("Noto Sans", "Noto Sans Devanagari", "Noto Sans Arabic"), size: 10.5pt)',
            '#set page(margin: 18mm)',
            '#set par(leading: 0.45em, spacing: 0.5em)',
            '',
            f'#text(size: 20pt, weight: "bold")[{_typst_escape(playlist_title or "YouTube Playlist Transcript")}]',
            '',
            f'#text(size: 10pt, fill: rgb("#787878"))[{len(videos)} videos • Generated by transcriptflow.io]',
        ]
        for i, v in enumerate(videos):
            title = _typst_escape(f"{i + 1}. {v.get('title') or v.get('video_id')}")
            words = v.get('word_count') if v.get('word_count') is not None else '?'
            meta = _typst_escape(
                f"Language: {v.get('language') or '?'}  •  Words: {words}  •  "
                f"Source: https://youtu.be/{v.get('video_id')}"
            )
            parts.append('#pagebreak()')
            parts.append(f'#text(size: 14pt, weight: "bold")[{title}]')
            parts.append('')
            parts.append(f'#text(size: 9pt, fill: rgb("#787878"))[{meta}]')
            parts.append('')
            # Each caption line as its own paragraph (blank-line separated) —
            # benchmarked smaller and faster than forced line breaks.
            body = '\n\n'.join(_typst_escape(line) for line in (v.get('transcript') or '').split('\n') if line.strip())
            parts.append(body)

        return _typst_compile('\n'.join(parts))
    except Exception as e:
        logger.warning(f"Typst combined-PDF build failed ({type(e).__name__}: {e}); falling back to fpdf2")
        return _build_combined_pdf_fpdf(playlist_title, videos)


def _build_combined_pdf_fpdf(playlist_title, videos):
    """Fallback combined-PDF builder using fpdf2 (slow for complex scripts at scale)."""
    from fpdf import FPDF

    text_sample = playlist_title + '\n' + '\n'.join(
        f"{v.get('title', '')}\n{v.get('transcript', '')}" for v in videos
    )
    pdf = FPDF()
    pdf.add_font('NotoSans', '', os.path.join(FONT_DIR, 'NotoSans-Regular.ttf'))
    pdf.add_font('NotoSans', 'B', os.path.join(FONT_DIR, 'NotoSans-Bold.ttf'))

    fallbacks = []
    if _DEVANAGARI_RE.search(text_sample):
        pdf.add_font('NotoDeva', '', os.path.join(FONT_DIR, 'NotoSansDevanagari-Regular.ttf'))
        fallbacks.append('NotoDeva')
    if _ARABIC_RE.search(text_sample):
        pdf.add_font('NotoArab', '', os.path.join(FONT_DIR, 'NotoSansArabic-Regular.ttf'))
        fallbacks.append('NotoArab')
    if fallbacks:
        pdf.set_fallback_fonts(fallbacks)
    if _COMPLEX_RE.search(text_sample):
        try:
            pdf.set_text_shaping(True)
        except Exception:
            pass

    pdf.set_auto_page_break(True, margin=18)

    pdf.add_page()
    pdf.set_font('NotoSans', 'B', 20)
    pdf.multi_cell(0, 10, playlist_title or 'YouTube Playlist Transcript', new_x='LMARGIN', new_y='NEXT')
    pdf.ln(2)
    pdf.set_font('NotoSans', '', 10)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(0, 6, f"{len(videos)} videos   •   Generated by transcriptflow.io", new_x='LMARGIN', new_y='NEXT')
    pdf.set_text_color(20, 20, 20)

    for i, v in enumerate(videos):
        pdf.add_page()
        pdf.set_font('NotoSans', 'B', 14)
        pdf.multi_cell(0, 8, f"{i + 1}. {v.get('title') or v.get('video_id')}", new_x='LMARGIN', new_y='NEXT')
        pdf.ln(1)
        pdf.set_font('NotoSans', '', 9)
        pdf.set_text_color(120, 120, 120)
        pdf.multi_cell(
            0, 5.5,
            f"Language: {v.get('language') or '?'}   •   Words: {v.get('word_count') if v.get('word_count') is not None else '?'}   •   "
            f"Source: https://youtu.be/{v.get('video_id')}",
            new_x='LMARGIN', new_y='NEXT'
        )
        pdf.ln(3)
        pdf.set_text_color(20, 20, 20)
        pdf.set_font('NotoSans', '', 10.5)
        for line in (v.get('transcript') or '').split('\n'):
            pdf.multi_cell(0, 6, line, new_x='LMARGIN', new_y='NEXT')

    return io.BytesIO(bytes(pdf.output()))


def _build_combined_docx(playlist_title, videos):
    """Build a single Word document containing every video's transcript, one section per video."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(playlist_title or 'YouTube Playlist Transcript', level=1)
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"{len(videos)} videos   •   Generated by transcriptflow.io")
    meta_run.font.size = Pt(9)

    for i, v in enumerate(videos):
        doc.add_page_break()
        doc.add_heading(f"{i + 1}. {v.get('title') or v.get('video_id')}", level=2)
        meta = doc.add_paragraph()
        meta_run = meta.add_run(
            f"Language: {v.get('language') or '?'}   •   Words: {v.get('word_count') if v.get('word_count') is not None else '?'}   •   "
            f"Source: https://youtu.be/{v.get('video_id')}"
        )
        meta_run.font.size = Pt(9)
        for line in (v.get('transcript') or '').split('\n'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            match = TIMESTAMP_LINE_RE.match(line)
            if match:
                ts_run = p.add_run(match.group(1) + ' ')
                ts_run.bold = True
                p.add_run(match.group(2))
            else:
                p.add_run(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf


COMBINED_EXPORT_FORMATS = {
    'pdf': ('application/pdf', _build_combined_pdf),
    'docx': ('application/vnd.openxmlformats-officedocument.wordprocessingml.document', _build_combined_docx),
}


# ---------------------------------------------------------------------------
# Playlist support (TranscriptFlow Pro)
# ---------------------------------------------------------------------------

PLAYLIST_ID_RE = re.compile(r'[?&]list=([a-zA-Z0-9_-]{10,})')
FREE_PLAYLIST_VIDEOS = int(os.environ.get('FREE_PLAYLIST_VIDEOS', 2))
PRO_MAX_PLAYLIST_VIDEOS = int(os.environ.get('PRO_MAX_PLAYLIST_VIDEOS', 100))
PRO_DAILY_VIDEO_QUOTA = int(os.environ.get('PRO_DAILY_VIDEO_QUOTA', 200))
PRO_MONTHLY_VIDEO_QUOTA = int(os.environ.get('PRO_MONTHLY_VIDEO_QUOTA', 1000))
DODO_API_BASE = os.environ.get('DODO_API_BASE', 'https://live.dodopayments.com')


def _playlist_id_from_url(url):
    match = PLAYLIST_ID_RE.search(url or '')
    return match.group(1) if match else None


def _fetch_playlist(playlist_id):
    """Fetch playlist metadata + first ~100 videos via YouTube's InnerTube API
    (through the proxy). Returns {'title': str, 'videos': [{video_id,title,duration}]}"""
    proxies = PROXY_CONFIG.to_requests_dict() if PROXY_CONFIG else None
    last_exc = None
    for attempt in range(3):
        try:
            resp = requests.post(
                'https://www.youtube.com/youtubei/v1/browse',
                json={
                    'context': {'client': {'clientName': 'WEB', 'clientVersion': '2.20240101.00.00'}},
                    'browseId': f'VL{playlist_id}',
                },
                proxies=proxies,
                timeout=20,
            )
            resp.raise_for_status()
            data = resp.json()
            break
        except Exception as e:
            last_exc = e
            logger.warning(f"Playlist fetch attempt {attempt + 1}/3 failed: {type(e).__name__}")
    else:
        raise last_exc

    title = None
    try:
        title = data['metadata']['playlistMetadataRenderer']['title']
    except Exception:
        try:
            title = data['header']['playlistHeaderRenderer']['title']['simpleText']
        except Exception:
            pass

    def _duration_to_seconds(text):
        try:
            parts = [int(p) for p in text.split(':')]
            seconds = 0
            for part in parts:
                seconds = seconds * 60 + part
            return seconds
        except Exception:
            return 0

    videos = []
    seen = set()

    def _collect(node):
        if len(videos) >= PRO_MAX_PLAYLIST_VIDEOS:
            return
        if isinstance(node, dict):
            # Current YouTube web structure (2025+): lockupViewModel entries
            lockup = node.get('lockupViewModel')
            if lockup and lockup.get('contentType') == 'LOCKUP_CONTENT_TYPE_VIDEO':
                video_id = lockup.get('contentId')
                if video_id and re.fullmatch(r'[a-zA-Z0-9_-]{11}', video_id) and video_id not in seen:
                    seen.add(video_id)
                    video_title = (lockup.get('metadata', {})
                                   .get('lockupMetadataViewModel', {})
                                   .get('title', {}) or {}).get('content') or video_id
                    duration_match = re.search(r'"text":\s*"(\d+(?::\d{2})+)"', json.dumps(lockup.get('contentImage', {})))
                    videos.append({
                        'video_id': video_id,
                        'title': video_title,
                        'duration_seconds': _duration_to_seconds(duration_match.group(1)) if duration_match else 0,
                        'playable': True,
                    })
                return
            # Legacy structure: playlistVideoRenderer entries
            renderer = node.get('playlistVideoRenderer')
            if renderer:
                video_id = renderer.get('videoId')
                if video_id and video_id not in seen:
                    seen.add(video_id)
                    try:
                        video_title = ''.join(run.get('text', '') for run in renderer['title'].get('runs', []))
                    except Exception:
                        video_title = video_id
                    videos.append({
                        'video_id': video_id,
                        'title': video_title,
                        'duration_seconds': int(renderer.get('lengthSeconds', 0) or 0),
                        'playable': renderer.get('isPlayable', True),
                    })
                return
            for value in node.values():
                _collect(value)
        elif isinstance(node, list):
            for item in node:
                _collect(item)

    _collect(data.get('contents', {}))

    if not videos:
        raise ValueError('empty playlist')
    return {'title': title or 'YouTube Playlist', 'videos': videos}


def _validate_license(license_key):
    """Validate a license key against Dodo Payments' public validation endpoint.
    Results are cached briefly to avoid hammering the API during bulk exports."""
    if not license_key or len(license_key) > 200:
        return False
    cache_key = f"license:{license_key}"
    cached = cache_get(cache_key)
    if cached is not None:
        return cached['valid']
    valid = False
    try:
        resp = requests.post(
            f'{DODO_API_BASE}/licenses/validate',
            json={'license_key': license_key},
            timeout=8,
        )
        if resp.ok:
            valid = bool(resp.json().get('valid'))
    except Exception as e:
        logger.warning(f"License validation error: {type(e).__name__}")
    _cache_short(cache_key, {'valid': valid}, ttl=600)
    return valid


# Short-TTL cache entries piggyback on the main LRU with their own expiry
def _cache_short(key, payload, ttl):
    payload = dict(payload)
    payload['_expires_at'] = time.time() + ttl
    cache_set(key, payload)


_orig_cache_get = cache_get


def cache_get(key):  # noqa: F811 - wrap to honor short TTLs
    entry = _orig_cache_get(key)
    if entry and '_expires_at' in entry and time.time() > entry['_expires_at']:
        return None
    return entry


# Per-license usage counters. Upstash Redis (free tier) when configured so
# quotas survive restarts; otherwise in-memory.
UPSTASH_URL = os.environ.get('UPSTASH_REDIS_REST_URL')
UPSTASH_TOKEN = os.environ.get('UPSTASH_REDIS_REST_TOKEN')
_quota_mem = {}
_quota_lock = threading.Lock()


def _quota_periods():
    now = time.gmtime()
    return time.strftime('%Y%m%d', now), time.strftime('%Y%m', now)


def _upstash(*command):
    resp = requests.post(
        f'{UPSTASH_URL}/pipeline',
        headers={'Authorization': f'Bearer {UPSTASH_TOKEN}'},
        json=[list(command)] if isinstance(command[0], str) else [list(c) for c in command],
        timeout=8,
    )
    resp.raise_for_status()
    return resp.json()


def check_and_consume_quota(license_key, n_videos):
    """Reserve n_videos against the license's daily/monthly quota.
    Returns (allowed: bool, info: dict)."""
    day, month = _quota_periods()
    if UPSTASH_URL and UPSTASH_TOKEN:
        try:
            result = _upstash(
                ['INCRBY', f'q:d:{day}:{license_key}', str(n_videos)],
                ['EXPIRE', f'q:d:{day}:{license_key}', '90000'],
                ['INCRBY', f'q:m:{month}:{license_key}', str(n_videos)],
                ['EXPIRE', f'q:m:{month}:{license_key}', '2764800'],
            )
            day_used = int(result[0]['result'])
            month_used = int(result[2]['result'])
        except Exception as e:
            logger.warning(f"Upstash quota error, falling back to memory: {type(e).__name__}")
            day_used = month_used = None
    else:
        day_used = month_used = None

    if day_used is None:
        with _quota_lock:
            rec = _quota_mem.setdefault(license_key, {})
            rec[f'd:{day}'] = rec.get(f'd:{day}', 0) + n_videos
            rec[f'm:{month}'] = rec.get(f'm:{month}', 0) + n_videos
            for k in [k for k in rec if not (k.endswith(day) or k.endswith(month))]:
                del rec[k]
            day_used, month_used = rec[f'd:{day}'], rec[f'm:{month}']

    allowed = day_used <= PRO_DAILY_VIDEO_QUOTA and month_used <= PRO_MONTHLY_VIDEO_QUOTA
    return allowed, {
        'day_used': day_used, 'day_limit': PRO_DAILY_VIDEO_QUOTA,
        'month_used': month_used, 'month_limit': PRO_MONTHLY_VIDEO_QUOTA,
    }


# Real transcript counter — the honest number shown on the site.
STATS_BASELINE = int(os.environ.get('STATS_BASELINE', 0))
_stats_mem = {'transcripts': 0}


def _increment_transcript_counter():
    if UPSTASH_URL and UPSTASH_TOKEN:
        try:
            _upstash(['INCR', 'stats:transcripts'])
            return
        except Exception:
            pass
    with _quota_lock:
        _stats_mem['transcripts'] += 1


def _get_transcript_count():
    count = None
    if UPSTASH_URL and UPSTASH_TOKEN:
        try:
            result = _upstash(['GET', 'stats:transcripts'])
            count = int(result[0]['result'] or 0)
        except Exception:
            pass
    if count is None:
        count = _stats_mem['transcripts']
    return STATS_BASELINE + count


@app.route('/api/stats', methods=['GET'])
def get_stats():
    """Public, real usage stats (no fabricated numbers on the site)."""
    return jsonify({'transcripts_total': _get_transcript_count(), 'success': True})


# ---------------------------------------------------------------------------
# On-site tool ratings (honest AggregateRating source — real votes only)
# ---------------------------------------------------------------------------

RATEABLE_PAGES = {'home', 'playlist', 'translate', 'download'}
_ratings_mem = {'seen': set(), 'sum': {}, 'cnt': {}}


def _rating_ip_hash():
    import hashlib
    ip = get_remote_address() or 'unknown'
    return hashlib.sha256(ip.encode()).hexdigest()[:24]


@app.route('/api/rate', methods=['POST'])
@limiter.limit('5 per minute')
def rate_tool():
    """Record a 1-5 star rating for a tool page. One vote per IP per page."""
    try:
        data = request.get_json() or {}
        page = data.get('page')
        rating = data.get('rating')
        if page not in RATEABLE_PAGES or not isinstance(rating, int) or not 1 <= rating <= 5:
            return jsonify({'error': 'Invalid rating.', 'error_type': 'invalid_input'}), 400

        ip_hash = _rating_ip_hash()
        counted = False
        if UPSTASH_URL and UPSTASH_TOKEN:
            try:
                result = _upstash(
                    ['SET', f'rate:seen:{page}:{ip_hash}', '1', 'NX', 'EX', '15552000'],  # 180 days
                )
                if result[0]['result'] == 'OK':
                    _upstash(
                        ['INCRBY', f'rate:sum:{page}', str(rating)],
                        ['INCR', f'rate:cnt:{page}'],
                    )
                    counted = True
            except Exception as e:
                logger.warning(f"Upstash rating error, falling back to memory: {type(e).__name__}")
                counted = None
        else:
            counted = None

        if counted is None:  # in-memory fallback
            with _quota_lock:
                key = f'{page}:{ip_hash}'
                if key not in _ratings_mem['seen']:
                    _ratings_mem['seen'].add(key)
                    _ratings_mem['sum'][page] = _ratings_mem['sum'].get(page, 0) + rating
                    _ratings_mem['cnt'][page] = _ratings_mem['cnt'].get(page, 0) + 1
                    counted = True
                else:
                    counted = False

        return jsonify({'success': True, 'counted': counted})
    except Exception as e:
        logger.error(f"Unexpected error in rate_tool: {str(e)}")
        return jsonify({'error': 'An unexpected error occurred.', 'error_type': 'server_error'}), 500


@app.route('/api/ratings', methods=['GET'])
@limiter.limit('30 per minute')
def get_ratings():
    """Live aggregate rating for a tool page — the number shown on the page IS the schema number."""
    page = request.args.get('page')
    if page not in RATEABLE_PAGES:
        return jsonify({'error': 'Invalid page.', 'error_type': 'invalid_input'}), 400

    cache_key = f'ratings:{page}'
    cached = cache_get(cache_key)
    if cached is not None:
        return jsonify(cached)

    total = count = 0
    if UPSTASH_URL and UPSTASH_TOKEN:
        try:
            result = _upstash(['GET', f'rate:sum:{page}'], ['GET', f'rate:cnt:{page}'])
            total = int(result[0]['result'] or 0)
            count = int(result[1]['result'] or 0)
        except Exception:
            total = count = -1
    else:
        total = count = -1

    if count == -1:  # in-memory fallback
        with _quota_lock:
            total = _ratings_mem['sum'].get(page, 0)
            count = _ratings_mem['cnt'].get(page, 0)

    payload = {
        'page': page,
        'average': round(total / count, 1) if count else 0,
        'count': count,
        'success': True,
    }
    _cache_short(cache_key, payload, ttl=60)
    return jsonify(payload)


@app.route('/api/playlist', methods=['POST'])
@limiter.limit('10 per minute; 60 per hour')
def get_playlist():
    """List a playlist's videos and which ones are unlocked for this user."""
    try:
        data = request.get_json() or {}
        playlist_id = _playlist_id_from_url(data.get('playlist_url'))
        if not playlist_id:
            return jsonify({
                'error': 'Invalid playlist URL. Paste a YouTube playlist link containing "list=".',
                'error_type': 'invalid_url'
            }), 400

        licensed = _validate_license((data.get('license_key') or '').strip())

        cache_key = f"playlist:{playlist_id}"
        playlist = cache_get(cache_key)
        if playlist is None:
            try:
                playlist = _fetch_playlist(playlist_id)
            except ValueError:
                return jsonify({
                    'error': 'This playlist appears to be empty, private, or unavailable.',
                    'error_type': 'playlist_unavailable'
                }), 404
            except Exception as e:
                logger.error(f"Playlist fetch failed for {playlist_id}: {type(e).__name__}")
                return jsonify({
                    'error': 'Could not load this playlist right now. Please try again.',
                    'error_type': 'fetch_error'
                }), 502
            cache_set(cache_key, playlist)

        videos = playlist['videos']
        playable_ids = [v['video_id'] for v in videos if v.get('playable', True)]
        unlocked = playable_ids if licensed else playable_ids[:FREE_PLAYLIST_VIDEOS]

        return jsonify({
            'playlist_id': playlist_id,
            'title': playlist['title'],
            'videos': videos,
            'video_count': len(videos),
            'unlocked_ids': unlocked,
            'licensed': licensed,
            'free_limit': FREE_PLAYLIST_VIDEOS,
            'success': True,
        })
    except Exception as e:
        logger.error(f"Unexpected error in get_playlist: {str(e)}")
        return jsonify({
            'error': 'An unexpected error occurred. Please try again later.',
            'error_type': 'server_error'
        }), 500


@app.route('/api/playlist/transcript', methods=['POST'])
@limiter.limit('60 per minute')
def playlist_transcript():
    """Fetch one transcript as part of a licensed playlist export (quota-metered).
    Free-tier videos go through the normal /api/transcript endpoint."""
    try:
        data = request.get_json() or {}
        video_id = _video_id_from_url(data.get('video_url')) or (
            data.get('video_id') if re.fullmatch(r'[a-zA-Z0-9_-]{11}', data.get('video_id') or '') else None)
        if not video_id:
            return jsonify({'error': 'Invalid video reference.', 'error_type': 'invalid_url'}), 400

        license_key = (data.get('license_key') or '').strip()
        if not _validate_license(license_key):
            return jsonify({
                'error': 'A valid TranscriptFlow Pro license is required for playlist exports.',
                'error_type': 'license_required'
            }), 402

        allowed, quota = check_and_consume_quota(license_key, 1)
        if not allowed:
            return jsonify({
                'error': 'You have reached your Pro export quota. It resets daily/monthly.',
                'error_type': 'quota_exceeded',
                'quota': quota,
            }), 429

        try:
            payload, was_cached = _transcript_payload(video_id, None)
        except Exception as e:
            error_response = _transcript_error_response(e, video_id)
            if error_response:
                return error_response
            logger.error(f"Playlist transcript fetch error: {str(e)}")
            return jsonify({'error': 'Failed to fetch transcript.', 'error_type': 'fetch_error'}), 500

        payload['cached'] = was_cached
        payload['quota'] = quota
        return jsonify(payload)
    except Exception as e:
        logger.error(f"Unexpected error in playlist_transcript: {str(e)}")
        return jsonify({
            'error': 'An unexpected error occurred. Please try again later.',
            'error_type': 'server_error'
        }), 500


@app.route('/api/playlist/export', methods=['POST'])
@limiter.limit('10 per minute; 30 per hour')
def export_playlist_combined():
    """Combine already-fetched playlist transcripts (from /api/playlist/transcript)
    into a single PDF or Word document. Formatting only — no YouTube/proxy calls,
    so this stays fast and cheap regardless of playlist size."""
    try:
        data = request.get_json() or {}
        export_format = (data.get('format') or '').lower().strip()
        if export_format not in COMBINED_EXPORT_FORMATS:
            return jsonify({
                'error': 'Invalid format. Supported formats: pdf, docx.',
                'error_type': 'invalid_input'
            }), 400

        videos = data.get('videos') or []
        if not isinstance(videos, list) or not videos:
            return jsonify({'error': 'No transcripts provided.', 'error_type': 'invalid_input'}), 400
        if len(videos) > PRO_MAX_PLAYLIST_VIDEOS:
            return jsonify({
                'error': f'Too many videos in one export (max {PRO_MAX_PLAYLIST_VIDEOS}).',
                'error_type': 'invalid_input'
            }), 400

        if len(videos) > FREE_PLAYLIST_VIDEOS:
            license_key = (data.get('license_key') or '').strip()
            if not _validate_license(license_key):
                return jsonify({
                    'error': f'A valid TranscriptFlow Pro license is required to combine more than {FREE_PLAYLIST_VIDEOS} videos.',
                    'error_type': 'license_required'
                }), 402

        # Per-video transcript length cap: real transcripts run ~10-100k chars even for
        # multi-hour videos; the cap stops abuse of this endpoint as a free PDF builder.
        max_chars = int(os.environ.get('COMBINED_EXPORT_MAX_CHARS_PER_VIDEO', 400_000))
        cleaned = []
        for v in videos[:PRO_MAX_PLAYLIST_VIDEOS]:
            if not isinstance(v, dict) or not isinstance(v.get('transcript'), str) or not v['transcript'].strip():
                continue
            cleaned.append({
                'video_id': re.sub(r'[^a-zA-Z0-9_-]', '', str(v.get('video_id') or ''))[:11],
                'title': str(v.get('title') or '')[:300],
                'transcript': v['transcript'][:max_chars],
                'language': str(v.get('language') or '')[:20],
                'word_count': v.get('word_count') if isinstance(v.get('word_count'), int) else None,
            })
        if not cleaned:
            return jsonify({'error': 'No valid transcripts provided.', 'error_type': 'invalid_input'}), 400

        playlist_title = str(data.get('playlist_title') or 'YouTube Playlist Transcript')[:300]
        mimetype, builder = COMBINED_EXPORT_FORMATS[export_format]
        buf = builder(playlist_title, cleaned)
        buf.seek(0)
        filename = f"{_slugify(playlist_title)}-combined.{export_format}"
        return send_file(buf, mimetype=mimetype, as_attachment=True, download_name=filename)

    except Exception as e:
        logger.error(f"Unexpected error in export_playlist_combined: {str(e)}")
        return jsonify({
            'error': 'An unexpected error occurred. Please try again later.',
            'error_type': 'server_error'
        }), 500


@app.route('/')
def index():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'message': 'TranscriptFlow Backend v3.4 is running',
        'version': '3.4.0',
        'proxy': 'configured' if PROXY_CONFIG else 'none',
        'timestamp': datetime.utcnow().isoformat()
    })

@app.route('/api/transcript', methods=['POST'])
@limiter.limit('10 per minute; 100 per hour')
def get_transcript():
    """Generate transcript from YouTube URL"""
    start_time = time.time()

    try:
        # Get request data
        data = request.get_json()
        if not data or not data.get('video_url'):
            return jsonify({
                'error': 'Invalid request. Please provide a video_url.',
                'error_type': 'invalid_input'
            }), 400

        video_url = data.get('video_url')
        target_language = (data.get('target_language') or '').strip() or None
        if target_language and not re.fullmatch(r'[a-zA-Z-]{2,12}', target_language):
            return jsonify({
                'error': 'Invalid target language code.',
                'error_type': 'invalid_input'
            }), 400

        # Extract video ID from URL. Only the ID is logged, never the full URL.
        video_id = _video_id_from_url(video_url)
        if not video_id:
            return jsonify({
                'error': 'Invalid YouTube URL format. Please provide a valid YouTube video URL.',
                'error_type': 'invalid_url'
            }), 400

        logger.info(f"Processing transcript request for video {video_id} (target_language={target_language})")

        try:
            payload, was_cached = _transcript_payload(video_id, target_language)
        except Exception as e:
            error_response = _transcript_error_response(e, video_id)
            if error_response:
                return error_response
            logger.error(f"Error fetching transcript: {str(e)}")
            return jsonify({
                'error': 'Failed to fetch transcript. Please try again later.',
                'error_type': 'fetch_error'
            }), 500

        payload['cached'] = was_cached
        payload['processing_time_ms'] = (time.time() - start_time) * 1000
        payload['timestamp'] = datetime.utcnow().isoformat()
        if was_cached:
            logger.info(f"Served video {video_id} from cache")
        return jsonify(payload)

    except Exception as e:
        logger.error(f"Unexpected error in get_transcript: {str(e)}")
        return jsonify({
            'error': 'An unexpected error occurred. Please try again later.',
            'error_type': 'server_error'
        }), 500


@app.route('/api/export', methods=['POST'])
@limiter.limit('10 per minute; 100 per hour')
def export_transcript():
    """Download the transcript as a TXT, Word (.docx), or PDF file."""
    try:
        data = request.get_json() or {}
        video_id = _video_id_from_url(data.get('video_url'))
        if not video_id:
            return jsonify({
                'error': 'Invalid YouTube URL format. Please provide a valid YouTube video URL.',
                'error_type': 'invalid_url'
            }), 400

        export_format = (data.get('format') or '').lower().strip()
        if export_format not in EXPORT_FORMATS:
            return jsonify({
                'error': 'Invalid format. Supported formats: txt, docx, pdf.',
                'error_type': 'invalid_input'
            }), 400

        target_language = (data.get('target_language') or '').strip() or None
        if target_language and not re.fullmatch(r'[a-zA-Z-]{2,12}', target_language):
            return jsonify({
                'error': 'Invalid target language code.',
                'error_type': 'invalid_input'
            }), 400

        logger.info(f"Export request for video {video_id} as {export_format}")

        try:
            payload, _ = _transcript_payload(video_id, target_language)
        except Exception as e:
            error_response = _transcript_error_response(e, video_id)
            if error_response:
                return error_response
            logger.error(f"Error fetching transcript for export: {str(e)}")
            return jsonify({
                'error': 'Failed to fetch transcript. Please try again later.',
                'error_type': 'fetch_error'
            }), 500

        mimetype, builder = EXPORT_FORMATS[export_format]
        filename = f"{_slugify(payload.get('video_title') or video_id)}.{export_format}"

        # Generated files are cached so repeat downloads are instant
        file_cache_key = f"export:{video_id}:{target_language or 'default'}:{export_format}"
        cached_file = cache_get(file_cache_key)
        if cached_file is not None:
            return send_file(io.BytesIO(cached_file['bytes']), mimetype=mimetype,
                             as_attachment=True, download_name=filename)

        if builder is None:
            buf = io.BytesIO((payload.get('transcript') or '').encode('utf-8'))
        else:
            buf = builder(payload)
        buf.seek(0)
        cache_set(file_cache_key, {'bytes': buf.getvalue()})
        return send_file(buf, mimetype=mimetype, as_attachment=True, download_name=filename)

    except Exception as e:
        logger.error(f"Unexpected error in export_transcript: {str(e)}")
        return jsonify({
            'error': 'An unexpected error occurred. Please try again later.',
            'error_type': 'server_error'
        }), 500


@app.route('/api/languages', methods=['POST'])
@limiter.limit('10 per minute; 100 per hour')
def get_languages():
    """List available transcript languages and translation targets for a video."""
    try:
        data = request.get_json()
        video_id = _video_id_from_url((data or {}).get('video_url'))
        if not video_id:
            return jsonify({
                'error': 'Invalid YouTube URL format. Please provide a valid YouTube video URL.',
                'error_type': 'invalid_url'
            }), 400

        cache_key = f"langs:{video_id}"
        cached = cache_get(cache_key)
        if cached is not None:
            return jsonify(cached)

        try:
            ytt_api = YouTubeTranscriptApi(proxy_config=PROXY_CONFIG)
            transcript_list = ytt_api.list(video_id)
            transcripts = [
                {
                    'language': t.language,
                    'language_code': t.language_code,
                    'is_generated': t.is_generated,
                }
                for t in transcript_list
            ]
            picked = _pick_transcript(ytt_api.list(video_id))
            translation_languages = [
                {'language': tl.language, 'language_code': tl.language_code}
                for tl in (picked.translation_languages if picked and picked.is_translatable else [])
            ]
        except (TranscriptsDisabled, VideoUnavailable, NoTranscriptFound):
            return jsonify({
                'error': 'No transcript available for this video.',
                'error_type': 'no_transcript'
            }), 404
        except (RequestBlocked, IpBlocked):
            return jsonify({
                'error': 'YouTube is temporarily blocking requests from our server. Please try again in a few minutes.',
                'error_type': 'ip_blocked'
            }), 503

        response_data = {
            'video_id': video_id,
            'transcripts': transcripts,
            'translation_languages': translation_languages,
            'success': True,
        }
        cache_set(cache_key, response_data)
        return jsonify(response_data)

    except Exception as e:
        logger.error(f"Unexpected error in get_languages: {str(e)}")
        return jsonify({
            'error': 'An unexpected error occurred. Please try again later.',
            'error_type': 'server_error'
        }), 500


@app.route('/api/health', methods=['GET'])
def health_check():
    """Detailed health check"""
    return jsonify({
        'status': 'healthy',
        'message': 'TranscriptFlow Backend v3.4 is running',
        'version': '3.4.0',
        'timestamp': datetime.utcnow().isoformat(),
        'endpoints': [
            '/api/transcript (POST)',
            '/api/health (GET)'
        ]
    })

# Error handlers
@app.errorhandler(429)
def rate_limit_handler(e):
    return jsonify({
        'error': 'Too many requests. Please wait a minute and try again.',
        'error_type': 'rate_limited'
    }), 429

@app.errorhandler(404)
def not_found_handler(e):
    return jsonify({
        'error': 'Endpoint not found',
        'error_type': 'not_found'
    }), 404

@app.errorhandler(500)
def internal_error_handler(e):
    logger.error(f"Internal server error: {str(e)}")
    return jsonify({
        'error': 'Internal server error',
        'error_type': 'internal_error'
    }), 500

if __name__ == '__main__':
    logger.info("Starting TranscriptFlow Backend v3.4")
    app.run(debug=False, host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
